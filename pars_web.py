from selenium import webdriver
import chromedriver_binary
import time
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup as bs
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx
import bankrot_fedres
import discval
import pravosudie
import kompromat_ru
import obreldoc

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink



def pars_web(inn):

        browser = webdriver.Chrome() #запускаем брайзер, может быть селениум Грид?
        browser.get('https://www.list-org.com/')
        browser.implicitly_wait(2)  # что бы сайт не закрывался
        time.sleep(1)

        find = browser.find_element(By.CLASS_NAME, 'autocomplete-input')
        find.send_keys(inn + "\n")
        time.sleep(1)

        result = browser.find_element(By.CSS_SELECTOR, '[class="card w-100 p-1 p-lg-3 mt-1"]')
        result = result.find_element(By.TAG_NAME, 'a')
        result.click()

        source_data = browser.page_source

        soup = bs(source_data, "lxml")  # отдаем данные в супчик

        title = soup.title.text
        maineinfo = soup.find("tbody").text
        for_komprom = title.split(',')

        browser.close()  # обяхательно закрываем сайт

        document = Document()

        document.add_heading(title, 0)
        document.add_paragraph('Общая информация', style='Intense Quote')
        p = document.add_paragraph(maineinfo)

        find_all_info_list_org = soup.find_all('div', class_="card w-100 p-1 p-lg-3 mt-2")

        for item in find_all_info_list_org:
            item_title = item.find('h6').find("div")
            item_content = item.find_all("div")
            document.add_paragraph(item_title, style='Intense Quote')
            document.add_paragraph(item_content[1].text)


        item_bankrot_fedres = bankrot_fedres.bankrot_fedres(inn)
        document.add_paragraph(item_bankrot_fedres[0], style='Intense Quote')
        document.add_paragraph('Результаты поиска: ' + item_bankrot_fedres[1])
        p = document.add_paragraph('Ссылка на ресурс: ')
        add_hyperlink(p, 'поиск по должникам на федресурсе', item_bankrot_fedres[2])
        document.add_paragraph('Время совершения запроса: ' + item_bankrot_fedres[3])

        item_pravosudie = pravosudie.pravosud(inn)
        document.add_paragraph(item_pravosudie[0], style='Intense Quote')
        document.add_paragraph('Результаты поиска: ' + item_pravosudie[1])
        p = document.add_paragraph('Ссылка на ресурс: ')
        add_hyperlink(p, 'поиск по делам и судебным актам', item_pravosudie[2])
        document.add_paragraph('Время совершения запроса: ' + item_pravosudie[3])

        item_discval = discval.obreldoc(inn)
        document.add_paragraph(item_discval[0], style='Intense Quote')
        document.add_paragraph('Результаты поиска: ' + item_discval[1])
        p = document.add_paragraph('Ссылка на ресурс: ')
        add_hyperlink(p, 'поиск по реестру дисквалифицированных лиц', item_discval[2])
        document.add_paragraph('Время совершения запроса: ' + item_discval[3])

        item_kompromat_ru = kompromat_ru.kompromat(for_komprom[0])
        document.add_paragraph(item_kompromat_ru[0], style='Intense Quote')
        document.add_paragraph('Результаты поиска: ' + item_kompromat_ru[1])
        p = document.add_paragraph('Ссылка на ресурс: ')
        add_hyperlink(p, 'поиск статей по сайту Компромат.ру', item_kompromat_ru[2])
        document.add_paragraph('Время совершения запроса: ' + item_kompromat_ru[3])

        item_obreldoc = obreldoc.obreldoc(inn)
        document.add_paragraph(item_obreldoc[0], style='Intense Quote')
        document.add_paragraph('Результаты поиска: ' + item_obreldoc[1])
        p = document.add_paragraph('Ссылка на ресурс: ')
        add_hyperlink(p, 'система информирования банков о состоянии обработки электронных документов', item_obreldoc[2])
        document.add_paragraph('Время совершения запроса: ' + item_obreldoc[3])

        name_of_file = for_komprom[0].replace('"', '') + '.docx'
        document.save(for_komprom[0].replace('"', '') + '.docx')

        lres = 'Парсинг успешно проведен, результат выслан в виде отчета.'

        totalres = [name_of_file, lres]
    #except:
        #totalres = 'Произошла ошибка при парсинге, повторите попытку через 5 секунд. ВНИМАНИЕ вы получили неверный отчет, он не должен был отправляться. Мы пока над эти работаем!'

        return totalres

